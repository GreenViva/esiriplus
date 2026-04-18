from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set default paragraph spacing
for s in doc.styles:
    try:
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.space_before = Pt(3)
    except:
        pass


def center_bold(text, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def heading(text, level=1, size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if level == 1:
        run.underline = True
    return p


def para(text, indent=False):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    return p


def numbered_para(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[{number:04d}] ")
    run.bold = True
    p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
center_bold("UNITED REPUBLIC OF TANZANIA", 14)
center_bold("BUSINESS REGISTRATIONS AND LICENSING AGENCY (BRELA)", 12)
center_bold("PATENTS ACT, 1987", 12)

doc.add_paragraph()
doc.add_paragraph()

center_bold("PATENT SPECIFICATION", 16)

doc.add_paragraph()

center_bold("Title of the Invention:", 12)
center_bold(
    "eSIRI PLUS: ANONYMOUS TELEMEDICINE PLATFORM WITH\n"
    "AUTOMATED NURSE-ASSISTED MEDICATION COMPLIANCE SYSTEM",
    14
)

doc.add_paragraph()
doc.add_paragraph()

# Applicant table
table = doc.add_table(rows=6, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ("Application No.:", "TZ/P/2026/______"),
    ("Filing Date:", datetime.datetime.now().strftime("%d/%m/%Y")),
    ("Applicant:", "PATRICK MOKIWA KIVANDA"),
    ("Nationality:", "Tanzanian"),
    ("Address:", "Kijoleri, Ilboru Ward, Arusha, Tanzania"),
    ("Inventor:", "PATRICK MOKIWA KIVANDA"),
]
for i, (label, value) in enumerate(data):
    cell_l = table.cell(i, 0)
    cell_r = table.cell(i, 1)
    run = cell_l.paragraphs[0].add_run(label)
    run.bold = True
    cell_r.paragraphs[0].add_run(value)

doc.add_paragraph()
doc.add_paragraph()
center_bold("International Patent Classification (IPC):", 11)
center_bold("G16H 10/60 — G16H 40/67 — H04L 9/32 — G06Q 20/00", 11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

center_bold("TABLE OF CONTENTS", 14)
doc.add_paragraph()

toc_items = [
    ("1.", "FIELD OF THE INVENTION", "3"),
    ("2.", "BACKGROUND OF THE INVENTION", "3"),
    ("3.", "SUMMARY OF THE INVENTION", "5"),
    ("4.", "BRIEF DESCRIPTION OF THE DRAWINGS", "7"),
    ("5.", "DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS", "8"),
    ("", "5.1 System Architecture Overview", "8"),
    ("", "5.2 Patient Anonymization and Identity System", "9"),
    ("", "5.3 Custom Authentication and Session Security", "10"),
    ("", "5.4 Multi-Tier Consultation Service Architecture", "11"),
    ("", "5.5 Consultation Request and Appointment Linking", "12"),
    ("", "5.6 Real-Time Session Management", "13"),
    ("", "5.7 Automated Nurse-Assisted Medication Compliance", "14"),
    ("", "5.8 Privacy-Preserving Notification System", "16"),
    ("", "5.9 AI-Assisted Report Generation", "17"),
    ("", "5.10 Mobile Payment Integration", "17"),
    ("", "5.11 Multi-Language Localization", "18"),
    ("6.", "CLAIMS", "19"),
    ("7.", "ABSTRACT", "22"),
    ("8.", "DRAWINGS AND FIGURES", "23"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    if num:
        run = p.add_run(f"{num}\t{title}")
        run.bold = True
    else:
        p.add_run(f"\t{title}")
    p.add_run(f"\t\t{page}")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. FIELD OF THE INVENTION
# ═══════════════════════════════════════════════════════════════════════════════

counter = [1]  # mutable counter for paragraph numbering


def np(text):
    """Numbered paragraph helper."""
    p = numbered_para(counter[0], text)
    counter[0] += 1
    return p


heading("1. FIELD OF THE INVENTION")

np(
    "The present invention relates generally to telemedicine systems and methods "
    "for delivering remote healthcare services. More specifically, it relates to "
    "an anonymous telemedicine platform that enables patients to receive medical "
    "consultations, prescriptions, and medication compliance monitoring without "
    "disclosing personally identifiable information."
)

np(
    "The invention further relates to automated nurse-assisted medication compliance "
    "systems, privacy-preserving authentication mechanisms for mobile health "
    "applications, real-time consultation session management, appointment-to-consultation "
    "linking, mobile payment integration for healthcare services, and AI-assisted "
    "medical report generation."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. BACKGROUND OF THE INVENTION
# ═══════════════════════════════════════════════════════════════════════════════

heading("2. BACKGROUND OF THE INVENTION")

np(
    "Telemedicine has emerged as a critical tool for healthcare delivery, "
    "particularly in resource-constrained environments where geographic barriers, "
    "limited infrastructure, and shortage of healthcare providers prevent timely "
    "access to medical services. Sub-Saharan Africa, and Tanzania in particular, "
    "faces acute healthcare access challenges with a doctor-to-patient ratio of "
    "approximately 1:25,000, compared to the WHO-recommended 1:1,000."
)

np(
    "Existing telemedicine platforms require patients to register with their full "
    "legal name, national identification number, email address, phone number, "
    "and residential address before accessing any services. This identity "
    "requirement creates a significant barrier in contexts where patients seek "
    "treatment for stigmatized conditions including sexually transmitted infections, "
    "mental health disorders, substance abuse, and reproductive health issues. "
    "Studies have documented that fear of identity disclosure causes up to 40% of "
    "patients with stigmatized conditions to avoid or delay seeking treatment."
)

np(
    "Furthermore, existing telemedicine systems do not address medication adherence "
    "after the consultation ends. The World Health Organization estimates that "
    "medication non-adherence in developing countries averages 50%, contributing to "
    "treatment failure, antimicrobial resistance, and preventable mortality. Current "
    "solutions are limited to automated text message or app-based reminders, which "
    "are passive and easily ignored. No existing telemedicine platform provides "
    "human-assisted medication compliance where a healthcare professional proactively "
    "contacts the patient to verify and encourage medication consumption."
)

np(
    "Prior art telemedicine systems also suffer from several additional limitations: "
    "(a) consultation sessions lack server-authoritative timing, allowing client-side "
    "clock manipulation to extend sessions without payment; (b) appointment systems "
    "exist independently from consultation systems, requiring manual coordination "
    "with no atomic linking; (c) push notification systems transmit sensitive health "
    "information in notification payloads, which may be cached, logged, or displayed "
    "on lock screens; (d) authentication systems rely on third-party identity providers "
    "that require email addresses or phone numbers, undermining anonymity."
)

np(
    "There is therefore a need for a telemedicine platform that: (i) enables fully "
    "anonymous patient access without collecting personally identifiable information; "
    "(ii) provides proactive human-assisted medication compliance monitoring; "
    "(iii) implements server-authoritative session management resistant to client "
    "manipulation; (iv) atomically links appointments to consultations; (v) delivers "
    "notifications without exposing health information in push payloads; and "
    "(vi) supports mobile money payment methods prevalent in the target market."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SUMMARY OF THE INVENTION
# ═══════════════════════════════════════════════════════════════════════════════

heading("3. SUMMARY OF THE INVENTION")

np(
    "The present invention provides an anonymous telemedicine platform and method, "
    "hereinafter referred to as 'eSIRI Plus', that addresses the aforementioned "
    "limitations of the prior art. The platform enables patients to access a full "
    "range of healthcare services—consultations, prescriptions, medication reminders, "
    "medical reports, and follow-up visits—without disclosing any personally "
    "identifiable information."
)

np(
    "In a first aspect, the invention provides a patient anonymization system that "
    "generates unique cryptographic patient identifiers (Patient IDs) following a "
    "structured format (e.g., 'ES-XXXX-XXXX') without requiring names, national "
    "identification numbers, email addresses, phone numbers, or residential addresses. "
    "The patient proves identity through user-defined security questions whose answers "
    "are hashed using PBKDF2, and sessions are bound to specific devices through "
    "cryptographic device binding."
)

np(
    "In a second aspect, the invention provides a multi-tier consultation service "
    "architecture with at least two tiers: an Economy tier providing a single follow-up "
    "consultation within a 14-day window, and a Royal tier providing unlimited follow-up "
    "consultations within a 14-day window plus access to nurse-assisted medication "
    "reminders. Six healthcare provider categories are supported: Nurse, Clinical "
    "Officer, Pharmacist, General Practitioner, Specialist, and Psychologist."
)

np(
    "In a third aspect, the invention provides an automated nurse-assisted medication "
    "compliance system. After a Royal-tier consultation, the prescribing doctor creates "
    "a medication consumption timetable specifying medication names, dosages, "
    "administration times, and duration. A server-side cron process detects approaching "
    "medication times, automatically assigns an available verified nurse from a pool, "
    "creates a voice-over-IP room, and dispatches the nurse to call the patient. "
    "The system includes retry logic for unavailable nurses, timeout-based reassignment, "
    "and text-based fallback notifications. Nurses receive compensation of 1,000 "
    "Tanzanian Shillings per completed reminder call."
)

np(
    "In a fourth aspect, the invention provides a custom HMAC-SHA256 JSON Web Token "
    "authentication system for patients that operates independently of third-party "
    "identity providers. The system features proactive token refresh, triple-layer "
    "token storage (encrypted storage, plain-text backup, in-memory cache), automatic "
    "retry with token refresh on 401 responses, and device binding verification."
)

np(
    "In a fifth aspect, the invention provides a consultation request system with "
    "60-second time-to-live requests, real-time status monitoring through WebSocket "
    "subscriptions with polling fallback, and appointment-to-consultation linking "
    "with one-appointment-one-consultation enforcement through unique constraints "
    "and optimistic locking."
)

np(
    "In a sixth aspect, the invention provides a privacy-preserving two-phase "
    "notification system where push notification payloads contain only opaque "
    "identifiers, and full content is fetched securely after device activation."
)

np(
    "In a seventh aspect, the invention provides an AI-assisted medical report "
    "generation system using large language models and QR-code-based report verification."
)

np(
    "In an eighth aspect, the invention provides mobile money payment integration "
    "using M-Pesa STK Push for consultation fees, service access, and video call "
    "recharges within the Tanzanian market."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. BRIEF DESCRIPTION OF THE DRAWINGS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("4. BRIEF DESCRIPTION OF THE DRAWINGS")

np(
    "The invention will be more fully understood from the following detailed "
    "description, taken in conjunction with the accompanying drawings, in which:"
)

figures = [
    ("Figure 1", "is a high-level system architecture diagram showing the client applications (Patient Android, Doctor Android), backend services (Supabase Edge Functions, PostgreSQL, Realtime), and external integrations (VideoSDK, Firebase Cloud Messaging, M-Pesa, OpenAI)."),
    ("Figure 2", "is a flow diagram illustrating the patient anonymization and session creation process, from anonymous ID generation through security question setup, device binding, and JWT token issuance."),
    ("Figure 3", "is a flow diagram illustrating the consultation request lifecycle, from patient request initiation through doctor notification, acceptance/rejection, and consultation session creation."),
    ("Figure 4", "is a flow diagram illustrating the appointment-to-consultation linking process, showing appointment validation, consultation request creation, doctor acceptance, and atomic appointment-consultation linkage with optimistic locking."),
    ("Figure 5", "is a sequence diagram illustrating the automated nurse-assisted medication compliance system, from medication timetable creation through cron detection, nurse assignment, voice call initiation, and completion callback with earnings credit."),
    ("Figure 6", "is a flow diagram illustrating the consultation session management system, showing server-authoritative timers, extension workflows, grace periods, payment processing, and automatic lifecycle transitions."),
    ("Figure 7", "is a diagram illustrating the privacy-preserving two-phase notification delivery system, showing the first phase (opaque payload push) and second phase (authenticated content fetch)."),
    ("Figure 8", "is a flow diagram illustrating the AI-assisted consultation report generation process, from doctor input through LLM processing to structured report output with QR verification code."),
    ("Figure 9", "is a diagram illustrating the triple-layer token storage and proactive refresh mechanism, showing encrypted storage, plain-text backup, in-memory cache, and the refresh decision flow."),
    ("Figure 10", "is a flow diagram illustrating the mobile payment integration for M-Pesa STK Push, showing payment initiation, STK prompt delivery, user confirmation, callback processing, and service activation."),
]

for fig_name, fig_desc in figures:
    p = doc.add_paragraph()
    run = p.add_run(fig_name + " ")
    run.bold = True
    p.add_run(fig_desc)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("5. DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS")

np(
    "The following detailed description sets forth specific embodiments of the "
    "invention. These embodiments are presented for illustrative purposes and are "
    "not intended to limit the scope of the invention as defined by the appended "
    "claims. It will be apparent to those skilled in the art that various modifications "
    "and variations can be made without departing from the spirit and scope of the "
    "invention."
)

# ── 5.1 System Architecture ──
heading("5.1 System Architecture Overview", level=2, size=12)

np(
    "Referring to Figure 1, the eSIRI Plus platform comprises the following principal "
    "components: (a) a Patient Client Application implemented as a native Android "
    "application using Kotlin and Jetpack Compose, structured as a multi-module Gradle "
    "project with core modules (common, domain, network, database) and feature modules "
    "(auth, patient, doctor, admin); (b) a Doctor Client Application sharing the same "
    "codebase but presenting doctor-specific feature modules; (c) a Backend Service "
    "Layer comprising Supabase Edge Functions implemented in TypeScript running on the "
    "Deno runtime; (d) a PostgreSQL Database with Row-Level Security (RLS) policies "
    "enforcing data isolation; (e) a Realtime Event Layer using Supabase Realtime "
    "WebSocket channels for live status updates; (f) an external Video Communication "
    "Service (VideoSDK) for voice and video calls; (g) Firebase Cloud Messaging (FCM) "
    "for push notification delivery; (h) M-Pesa payment gateway for mobile money "
    "transactions; and (i) an AI Service (OpenAI) for medical report generation."
)

np(
    "The client applications communicate with the backend exclusively through HTTPS "
    "REST API calls to Edge Functions and WebSocket connections for real-time events. "
    "The network layer uses OkHttp with Retrofit for REST calls and the Supabase-kt "
    "library with the OkHttp Ktor engine for Edge Function invocations and Realtime "
    "subscriptions. An interceptor chain comprising ProactiveRefresh, Auth, Retry, "
    "and Logging interceptors processes all outgoing requests. A TokenRefreshAuthenticator "
    "handles 401 responses with automatic token refresh and request retry."
)

np(
    "Local data persistence on the client device uses Room 2.7.1 with SQLCipher "
    "encryption (AES-256). The database encryption key is derived from the Android "
    "Keystore and stored in EncryptedSharedPreferences. The database contains 29 "
    "entities covering users, sessions, consultations, payments, messages, notifications, "
    "prescriptions, diagnoses, vital signs, and audit logs. Backup rules explicitly "
    "exclude the encrypted database and passphrase from cloud and device backups."
)

# ── 5.2 Patient Anonymization ──
heading("5.2 Patient Anonymization and Identity System", level=2, size=12)

np(
    "Referring to Figure 2, the patient anonymization system operates as follows. "
    "When a new patient accesses the platform for the first time, the system generates "
    "a unique Patient ID following a structured format comprising a two-letter prefix "
    "'ES' followed by two groups of four alphanumeric characters separated by hyphens "
    "(e.g., 'ES-A7K2-M9P4'). No personally identifiable information is collected—the "
    "patient is not asked for their name, national ID number, email address, phone "
    "number, or residential address at any point."
)

np(
    "The patient session creation process invokes a server-side Edge Function "
    "('create-patient-session') that generates: (i) a universally unique session "
    "identifier (UUID v4); (ii) a cryptographically random session token of at least "
    "64 characters using the Web Crypto API; (iii) a separate refresh token of at "
    "least 96 characters. The session token is hashed using both SHA-256 (for fast "
    "indexed database lookups) and bcrypt with a cost factor of 10 (for timing-attack "
    "resistant verification). Both hashes are stored in the patient_sessions table."
)

np(
    "The server issues a custom JSON Web Token signed using HMAC-SHA256 with a "
    "server-side secret (JWT_SECRET environment variable). The token payload contains: "
    "sub (session identifier), session_id, role set to 'authenticated' for PostgreSQL "
    "RLS compatibility, app_role set to 'patient', and standard exp/iat timestamps. "
    "The access token has a configurable expiration (default: 1 hour); the refresh "
    "token has a longer expiration (default: 30 days)."
)

np(
    "Account recovery is achieved through five user-defined security questions. The "
    "patient selects five questions from a predefined set and provides answers. Each "
    "answer is normalized (trimmed, lowercased) and hashed using PBKDF2 with a unique "
    "salt before storage. Recovery requires correctly answering at least four of the "
    "five questions. An account lockout mechanism engages after a configurable number "
    "of failed recovery attempts (default: 5), with progressive lockout durations."
)

np(
    "Device binding cryptographically links the patient session to a specific device "
    "using the 'check-device-binding' Edge Function. The device's unique identifier "
    "is hashed and stored alongside the session. Subsequent requests include the "
    "device identifier for verification, preventing unauthorized session use on "
    "different devices."
)

# ── 5.3 Authentication ──
heading("5.3 Custom Authentication and Session Security", level=2, size=12)

np(
    "Referring to Figure 9, the authentication system implements a triple-layer token "
    "storage mechanism on the client device: (a) Primary Layer: EncryptedSharedPreferences "
    "using AES-256-GCM encryption, backed by the Android Keystore; (b) Backup Layer: "
    "standard SharedPreferences as a fallback when encrypted storage fails due to "
    "Keystore corruption or device-specific issues; (c) In-Memory Layer: a volatile "
    "in-memory cache in the TokenManager singleton providing instant access without "
    "I/O operations."
)

np(
    "Token refresh follows a proactive strategy: before each API call, the client "
    "checks whether the access token will expire within a configurable window "
    "(default: 5 minutes). If so, the client invokes the 'refresh-patient-session' "
    "Edge Function with the refresh token to obtain new access and refresh tokens "
    "before proceeding with the original request. If the proactive refresh fails but "
    "a non-expired token is still available, the request proceeds with the existing "
    "token."
)

np(
    "Upon receiving a 401 Unauthorized response, the TokenRefreshAuthenticator "
    "intercepts the response, attempts a single token refresh, and retries the "
    "original request with the new token. If the refresh also fails, the user is "
    "presented with a session expiration message. A guard prevents concurrent refresh "
    "attempts from multiple threads using a ReentrantLock."
)

np(
    "The server validates every request by: (i) extracting the Bearer token from the "
    "Authorization header; (ii) verifying the HMAC-SHA256 signature against the "
    "JWT_SECRET; (iii) checking the token has not expired; (iv) verifying the app_role "
    "claim matches the expected role for the endpoint; (v) querying the patient_sessions "
    "table to confirm the session exists and is not revoked."
)

# ── 5.4 Multi-Tier ──
heading("5.4 Multi-Tier Consultation Service Architecture", level=2, size=12)

np(
    "The platform implements a tiered service model with differentiated access levels. "
    "In the preferred embodiment, two tiers are provided: (a) Economy Tier: provides a "
    "single consultation session with one follow-up consultation permitted within a "
    "14-day window from the original consultation date; (b) Royal Tier: provides a "
    "consultation session with unlimited follow-up consultations within a 14-day "
    "window, plus access to the automated nurse-assisted medication compliance system."
)

np(
    "Service tiers are pre-populated in the database during initialization with "
    "tier-specific pricing in Tanzanian Shillings (TZS). Six healthcare provider "
    "categories are supported, each with configurable consultation duration and pricing: "
    "Nurse, Clinical Officer, Pharmacist, General Practitioner, Specialist, and "
    "Psychologist. The consultation duration is determined by a server-side function "
    "based on the provider's specialty (e.g., 15 minutes for Nurse, 30 minutes for "
    "Specialist, 45 minutes for Psychologist)."
)

np(
    "Follow-up consultations are tracked per consultation record. The follow-up system "
    "checks: (i) the original consultation was completed (status = 'completed' and "
    "report_submitted = true); (ii) the current date is within the 14-day follow-up "
    "window; (iii) for Economy tier, the follow-up count has not exceeded the maximum "
    "of 1; (iv) for Royal tier, no follow-up count limit is enforced."
)

# ── 5.5 Consultation Request ──
heading("5.5 Consultation Request and Appointment Linking", level=2, size=12)

np(
    "Referring to Figures 3 and 4, the consultation request system operates with a "
    "60-second time-to-live (TTL) mechanism. When a patient initiates a consultation "
    "request, the client application: (i) proactively refreshes the authentication "
    "token if expiring within 5 minutes; (ii) invokes the 'handle-consultation-request' "
    "Edge Function with action 'create' and the request parameters; (iii) upon "
    "successful creation, starts a local 60-second countdown timer; (iv) subscribes "
    "to Supabase Realtime WebSocket channel filtered by the patient's session ID; "
    "(v) simultaneously starts a 3-second polling loop as a fallback for environments "
    "where WebSocket connectivity is unreliable."
)

np(
    "The server-side Edge Function validates the request, checks the target doctor's "
    "availability (is_available = true, in_session = false, is_verified = true), "
    "and creates a consultation_requests record with status 'pending'. If the "
    "doctor is unavailable or already in a session, the server returns an error "
    "with a specific code ('doctor_in_session' or 'not currently available')."
)

np(
    "When the request includes an appointment_id, additional validation is performed: "
    "(i) the appointment exists and belongs to the requesting patient; (ii) the "
    "appointment has no existing linked consultation (consultation_id IS NULL); "
    "(iii) the appointment status permits consultation initiation (BOOKED, CONFIRMED, "
    "or MISSED); (iv) no pending consultation request exists for the same appointment. "
    "Upon doctor acceptance, the system atomically updates the appointment record using "
    "an optimistic lock condition (WHERE consultation_id IS NULL) to link the newly "
    "created consultation and transitions the appointment status to IN_PROGRESS."
)

np(
    "If the 60-second TTL expires without doctor response, the client invokes the "
    "'expire' action on the Edge Function to transition the request status. A "
    "server-side cron job ('expire-followup-escrow') independently expires stale "
    "requests as a safety net."
)

# ── 5.6 Session Management ──
heading("5.6 Real-Time Session Management", level=2, size=12)

np(
    "Referring to Figure 6, consultation sessions are managed with server-authoritative "
    "timing. When a doctor accepts a consultation request, the server: (i) creates a "
    "consultation record with status 'ACTIVE'; (ii) sets the doctor's in_session flag "
    "to true via a database trigger; (iii) calculates the scheduled end time based on "
    "the service duration function; (iv) returns the session details to both client "
    "applications via Realtime channels."
)

np(
    "Client-side timer synchronization uses the server-provided scheduled_end_time and "
    "server_time to compute a remaining duration. The client anchors this duration to "
    "a monotonic clock (SystemClock.elapsedRealtime() on Android), preventing "
    "manipulation through system clock changes. The timer ticks locally but is "
    "periodically re-synchronized against the server time."
)

np(
    "Session extension follows a multi-step workflow: (i) the doctor initiates an "
    "extension request, transitioning the consultation to 'AWAITING_EXTENSION' status; "
    "(ii) the patient receives a real-time notification and can accept or decline; "
    "(iii) upon acceptance, a 'GRACE_PERIOD' status is entered with a configurable "
    "duration (default: 3 minutes) during which the patient processes payment via "
    "M-Pesa STK Push; (iv) upon payment confirmation (verified via callback), the "
    "consultation resumes with the extended duration added."
)

np(
    "Consultation completion follows a defined sequence: (i) the doctor ends the "
    "consultation, transitioning to 'COMPLETED_REPORT_PENDING' status; (ii) the "
    "doctor submits a consultation report (see Section 5.9); (iii) the report "
    "submission trigger sets report_submitted = true, which fires a database trigger "
    "to set in_session = false, making the doctor available for new requests."
)

# ── 5.7 Medication Compliance ──
doc.add_page_break()
heading("5.7 Automated Nurse-Assisted Medication Compliance System", level=2, size=12)

np(
    "Referring to Figure 5, the medication compliance system is a novel feature "
    "exclusive to Royal-tier consultations. The system comprises the following "
    "components and workflow:"
)

np(
    "TIMETABLE CREATION: During consultation report submission for a Royal-tier "
    "patient, the prescribing doctor creates one or more medication consumption "
    "timetables through the client application. Each timetable specifies: "
    "medication_name (string), dosage (string, e.g., '500mg'), pharmaceutical_form "
    "(string, e.g., 'tablet'), times_per_day (integer, 1-4), scheduled_times "
    "(array of time strings in HH:MM format, e.g., ['08:00', '14:00', '20:00']), "
    "and duration_days (integer). Smart defaults are provided: 1x daily = 08:00; "
    "2x daily = 08:00, 20:00; 3x daily = 08:00, 14:00, 20:00; 4x daily = 06:00, "
    "12:00, 18:00, 22:00."
)

np(
    "The timetable is stored in the medication_timetables table with fields: "
    "timetable_id (UUID, primary key), consultation_id (foreign key), "
    "patient_session_id (foreign key), doctor_id (foreign key), medication_name, "
    "dosage, form, times_per_day, scheduled_times (TEXT array), duration_days, "
    "start_date, end_date (computed as start_date + duration_days), is_active "
    "(boolean, default true), and timestamps. Row-Level Security policies ensure "
    "doctors can only access timetables they created, patients can only see their "
    "own timetables, and the service role has full access."
)

np(
    "CRON DETECTION: A pg_cron scheduled job fires every minute, invoking the "
    "'medication-reminder-cron' Edge Function via HTTP POST. The function: "
    "(i) validates the request using a shared X-Cron-Secret header; "
    "(ii) converts the current UTC server time to East Africa Time (EAT, UTC+3); "
    "(iii) extracts the current time as HH:MM and current date; "
    "(iv) queries active timetables where the current time matches ANY element "
    "in the scheduled_times array AND the current date falls between start_date "
    "and end_date (inclusive)."
)

np(
    "EVENT CREATION AND DEDUPLICATION: For each matching timetable, the system "
    "inserts a record into medication_reminder_events with a UNIQUE constraint on "
    "(timetable_id, scheduled_date, scheduled_time). The INSERT uses ON CONFLICT "
    "DO NOTHING to prevent duplicate events if the cron fires multiple times "
    "within the same minute window. Each event is created with status 'pending' "
    "and retry_count = 0."
)

np(
    "NURSE ASSIGNMENT: For each new pending event, the system queries the "
    "doctor_profiles table (which stores all healthcare providers including nurses) "
    "filtered by: specialty = 'nurse', is_verified = true, is_available = true, "
    "in_session = false. The first matching nurse is selected. If no nurse is "
    "available, the event status is set to 'no_nurse' for retry on subsequent "
    "cron executions."
)

np(
    "VOICE CALL INITIATION: Upon identifying an available nurse, the system: "
    "(i) creates a voice-over-IP room via the VideoSDK REST API (configured for "
    "audio-only communication); (ii) updates the event with the nurse_id, "
    "video_room_id, status = 'nurse_notified', and notified_at timestamp; "
    "(iii) dispatches a high-priority FCM data message to the nurse's device with "
    "type 'MEDICATION_REMINDER_CALL', containing the room_id, patient_session_id, "
    "medication_name, and dosage; (iv) dispatches a separate notification to the "
    "patient with type 'MEDICATION_REMINDER_PATIENT' indicating that a nurse will "
    "call shortly."
)

np(
    "NURSE CALL HANDLING: On the nurse's Android device, the FCM message triggers "
    "the IncomingCallService, which acquires a partial wake lock and displays a "
    "full-screen incoming call notification with caller role 'medication_reminder' "
    "and audio-only mode. The notification displays 'Medication Reminder: [medication "
    "name]'. The nurse accepts the call, which joins the VideoSDK room. The nurse "
    "speaks with the patient, confirms medication consumption, and ends the call."
)

np(
    "COMPLETION CALLBACK: After the call, the nurse's client application invokes "
    "the 'medication-reminder-callback' Edge Function with action 'completed' and "
    "the event_id. The server: (i) updates the event status to 'completed' and "
    "records call_ended_at; (ii) credits the nurse with 1,000 TZS by inserting a "
    "record into doctor_earnings with earning_type = 'medication_reminder', "
    "amount = 1000, and currency = 'TZS'."
)

np(
    "RETRY AND FALLBACK MECHANISMS: The cron function implements three recovery "
    "mechanisms on each execution: (a) Retry: events with status 'no_nurse' and "
    "retry_count < 3 are re-processed, incrementing the retry_count; (b) Timeout: "
    "events with status 'nurse_notified' where the notified_at timestamp is older "
    "than 2 minutes are reassigned to a different nurse (excluding the previously "
    "assigned nurse), up to a maximum of 2 reassignments; (c) Fallback: after "
    "exhausting all retry and reassignment attempts, the system sends a text-only "
    "push notification directly to the patient with the medication name and dosage "
    "as a passive reminder."
)

# ── 5.8 Notification ──
heading("5.8 Privacy-Preserving Notification System", level=2, size=12)

np(
    "Referring to Figure 7, the notification system implements a two-phase delivery "
    "mechanism designed to prevent exposure of sensitive health information through "
    "push notification payloads. In Phase 1, the server sends an FCM message "
    "containing only: (i) a notification_id (opaque UUID); (ii) a type string "
    "(e.g., 'CONSULTATION_REQUEST', 'MESSAGE', 'PAYMENT_CONFIRMED'); and "
    "(iii) a session_id for routing. No patient names, doctor names, medication "
    "names, consultation details, or health information is included in the push "
    "payload."
)

np(
    "In Phase 2, upon device activation from the push notification, the client "
    "application: (i) immediately displays a placeholder notification with a "
    "type-appropriate generic title (e.g., 'New Message' for message type); "
    "(ii) authenticates with the server using stored credentials; (iii) fetches "
    "the full notification content from the server; (iv) updates the notification "
    "display with the actual title and body. If the content fetch fails, the "
    "generic notification remains visible."
)

np(
    "For incoming call notifications (types 'VIDEO_CALL_INCOMING' and "
    "'MEDICATION_REMINDER_CALL'), the push payload is configured as a data-only "
    "message without a notification display block. This enables the "
    "EsiriplusFirebaseMessagingService to start a foreground IncomingCallService "
    "that acquires a partial wake lock and displays a full-screen call notification, "
    "which is critical for devices with aggressive background process management "
    "(e.g., OnePlus, Xiaomi, Samsung)."
)

# ── 5.9 AI Report ──
heading("5.9 AI-Assisted Report Generation", level=2, size=12)

np(
    "Referring to Figure 8, the report generation system operates through the "
    "'generate-consultation-report' Edge Function. The doctor submits: consultation_id, "
    "diagnosed_problem, medical_category, severity (mild/moderate/severe/critical), "
    "treatment_plan text, and optional prescriptions (each with medication_name, form, "
    "dosage, instructions, route). The server fetches the consultation's service_tier "
    "from the database."
)

np(
    "The function constructs a prompt for a large language model (OpenAI GPT-4) "
    "instructing it to generate five structured text fields: (i) presenting_symptoms "
    "(synthesized from the diagnosed problem and consultation context); "
    "(ii) diagnosis_assessment (clinical assessment prose); (iii) treatment_plan "
    "(expanded treatment plan incorporating medications); (iv) prescribed_medications "
    "(formal dosage instructions); (v) follow_up_instructions. A verification_code "
    "(UUID v4) is generated and embedded in the report for authenticity verification "
    "accessible via QR code scanning."
)

np(
    "For Royal-tier consultations, the report submission endpoint additionally accepts "
    "an optional medication_timetables array. When present, the server creates "
    "corresponding medication_timetables records for the nurse-assisted compliance "
    "system described in Section 5.7."
)

# ── 5.10 Payment ──
heading("5.10 Mobile Payment Integration", level=2, size=12)

np(
    "Referring to Figure 10, the platform integrates with the M-Pesa mobile money "
    "system prevalent in Tanzania. Three payment types are supported: (i) service "
    "access payments for purchasing consultation tier access; (ii) consultation fee "
    "payments during session extension grace periods; (iii) video call recharge "
    "payments for purchasing voice/video call minutes."
)

np(
    "The payment flow operates as follows: (a) the client sends a payment request "
    "to the appropriate Edge Function; (b) the server initiates an M-Pesa STK Push "
    "request, which delivers a USSD prompt to the patient's mobile phone; (c) the "
    "patient enters their M-Pesa PIN on their phone to authorize the payment; "
    "(d) M-Pesa processes the transaction and sends a callback to the server's "
    "'mpesa-callback' Edge Function; (e) the server verifies the callback, updates "
    "the payment status, and activates the corresponding service."
)

# ── 5.11 Localization ──
heading("5.11 Multi-Language Localization", level=2, size=12)

np(
    "The platform supports six languages: English (default), Swahili (Kiswahili), "
    "Arabic, French, Spanish, and Hindi. Localization is implemented through Android "
    "resource string files (strings.xml) with language-qualified directories "
    "(values-sw/, values-ar/, values-fr/, values-es/, values-hi/). All user-facing "
    "strings across the patient, doctor, and app modules are localized. Right-to-left "
    "layout support is provided for Arabic."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. CLAIMS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("6. CLAIMS")

np(
    "What is claimed is:"
)

claims = [
    # CLAIM 1
    (
        "An anonymous telemedicine platform system, comprising:\n"
        "a. a patient anonymization module that generates unique cryptographic patient identifiers "
        "without collecting personally identifiable information, wherein each patient identifier "
        "follows a format of a prefix code followed by alphanumeric segments separated by hyphens;\n"
        "b. a multi-tier consultation service module providing at least two service tiers with "
        "differentiated follow-up access, wherein a first tier provides a single follow-up "
        "consultation and a second tier provides unlimited follow-up consultations within a "
        "configurable time window;\n"
        "c. a consultation request module that creates time-limited consultation requests directed "
        "to specific healthcare providers, monitors request status through real-time WebSocket "
        "subscriptions and periodic polling fallback, and links consultation requests to "
        "pre-existing appointments;\n"
        "d. a session management module that maintains server-authoritative consultation timers "
        "with extension workflows, grace periods, and automatic lifecycle transitions;\n"
        "e. a custom authentication module using cryptographic JSON Web Tokens signed with "
        "HMAC-SHA256 for patient sessions, with proactive token refresh, triple-layer token "
        "storage fallback comprising encrypted storage, plain-text backup storage, and in-memory "
        "cache, and automatic retry with token refresh on authentication failures;\n"
        "f. a privacy-preserving notification module wherein push notification payloads contain "
        "only opaque identifiers and notification types, and full notification content is fetched "
        "securely from a server after device activation;\n"
        "g. a mobile payment integration module supporting STK Push-based mobile money transactions "
        "for consultation fees and service access;\n"
        "h. an AI-assisted report generation module that processes consultation data through a "
        "large language model to produce structured medical reports with verification codes; and\n"
        "i. a multi-language localization module supporting at least six languages."
    ),
    # CLAIM 2
    (
        "The anonymous telemedicine platform system as claimed in claim 1, further comprising an "
        "automated nurse-assisted medication compliance system, wherein:\n"
        "a. a prescribing healthcare provider creates a medication consumption timetable specifying "
        "at least a medication name, number of daily administrations, specific administration times, "
        "and a duration in days;\n"
        "b. the timetable is stored in a server database with a unique timetable identifier, "
        "consultation reference, patient session identifier, medication details, an array of "
        "scheduled administration times, duration, computed end date, and active status flag;\n"
        "c. a server-side scheduled process executes at a configurable interval, converts current "
        "server time to a local timezone, and queries active timetables where the current time "
        "matches any entry in the scheduled times array;\n"
        "d. for each matching entry, the system inserts a reminder event record with a unique "
        "constraint preventing duplicate processing;\n"
        "e. the system queries available healthcare providers filtered by specialty, verification "
        "status, and availability to identify an available nurse;\n"
        "f. upon identifying a nurse, the system creates a voice-over-IP room, dispatches a "
        "call notification to the nurse, and a pending-call notification to the patient;\n"
        "g. the nurse receives, accepts, and conducts the voice call with the patient; and\n"
        "h. upon completion, the system credits the nurse with a predetermined compensation amount."
    ),
    # CLAIM 3
    (
        "The system as claimed in claim 2, wherein the medication compliance system further "
        "comprises:\n"
        "a. a retry mechanism re-processing events where no nurse was available, up to a "
        "configurable maximum retry count;\n"
        "b. a timeout mechanism reassigning events where the notified nurse has not initiated "
        "a call within a configurable timeout, excluding previously assigned nurses;\n"
        "c. a patient fallback mechanism sending a text-only notification to the patient after "
        "exhausting all retry and reassignment attempts; and\n"
        "d. a patient-unreachable mechanism enabling the nurse to report that the patient did "
        "not answer, triggering a missed-reminder notification."
    ),
    # CLAIM 4
    (
        "The system as claimed in claim 1, wherein the consultation request module further "
        "comprises an appointment linking mechanism, wherein:\n"
        "a. a consultation request includes an optional appointment identifier;\n"
        "b. the system validates that the appointment exists, belongs to the requesting patient, "
        "has no existing linked consultation, and has a permissible status;\n"
        "c. the system validates that no pending request exists for the same appointment;\n"
        "d. upon provider acceptance, the system atomically links the consultation to the "
        "appointment using an optimistic lock condition; and\n"
        "e. appointments remain in the active list until a consultation is successfully linked."
    ),
    # CLAIM 5
    (
        "The system as claimed in claim 1, wherein the patient anonymization module comprises:\n"
        "a. session creation generating a UUID session identifier, a random session token of "
        "at least 64 characters, and a refresh token of at least 96 characters;\n"
        "b. a custom JWT signed with HMAC-SHA256 containing session identifier, role claims "
        "for database row-level security compatibility, and expiration timestamps;\n"
        "c. dual hashing of session tokens using SHA-256 for indexed lookups and bcrypt for "
        "timing-attack resistant verification;\n"
        "d. account recovery through five security questions with PBKDF2-hashed answers and "
        "progressive lockout after failed attempts; and\n"
        "e. device binding cryptographically linking sessions to specific devices."
    ),
    # CLAIM 6
    (
        "The system as claimed in claim 1, wherein the session management module comprises:\n"
        "a. a server-authoritative timer with duration determined by provider specialty;\n"
        "b. client timer synchronization using monotonic clock anchoring against server time;\n"
        "c. an extension workflow with provider initiation, patient acceptance, a configurable "
        "grace period for payment processing, and automatic session resumption upon payment;\n"
        "d. automatic lifecycle management via database triggers synchronizing provider "
        "availability based on consultation state; and\n"
        "e. a follow-up system with tier-configurable maximum follow-up counts within a "
        "configurable time window."
    ),
    # CLAIM 7
    (
        "The system as claimed in claim 1, wherein the privacy-preserving notification module "
        "comprises:\n"
        "a. a two-phase delivery system where the first phase transmits only opaque identifiers "
        "and type strings, with no health information in push payloads;\n"
        "b. a second phase wherein the client authenticates and fetches full content from the "
        "server after device activation;\n"
        "c. generic placeholder notifications displayed immediately while content fetch is "
        "in progress;\n"
        "d. data-only FCM messages for incoming call notifications, enabling foreground service "
        "wake lock acquisition and full-screen call display; and\n"
        "e. FCM token synchronization with separate storage for provider and patient accounts."
    ),
    # CLAIM 8
    (
        "The system as claimed in claim 1, wherein the AI-assisted report generation module "
        "comprises:\n"
        "a. a report input interface accepting consultation identifier, diagnosed problem, "
        "medical category, severity, treatment plan, and optional prescriptions;\n"
        "b. large language model integration generating five structured output fields;\n"
        "c. a UUID verification code embedded in the report for authenticity verification;\n"
        "d. a report submission trigger that clears provider in-session status; and\n"
        "e. for premium-tier consultations, optional medication timetable creation for the "
        "automated nurse-assisted compliance system."
    ),
    # CLAIM 9
    (
        "A method for providing anonymous telemedicine services with automated medication "
        "compliance, comprising:\n"
        "a. generating a unique patient identifier without collecting personally identifiable "
        "information;\n"
        "b. establishing a session using HMAC-SHA256 signed JWT with triple-layer token storage;\n"
        "c. receiving a consultation request optionally linked to a pre-existing appointment;\n"
        "d. monitoring request status through concurrent WebSocket and polling mechanisms;\n"
        "e. upon acceptance, creating a session with server-authoritative timing;\n"
        "f. managing the session lifecycle through extension, grace period, and payment workflows;\n"
        "g. generating an AI-assisted medical report upon completion;\n"
        "h. for premium-tier consultations, receiving and storing medication timetables;\n"
        "i. executing a scheduled process to detect medication administration times;\n"
        "j. assigning available nurses, creating voice channels, and dispatching notifications;\n"
        "k. crediting nurse compensation upon call completion; and\n"
        "l. delivering all notifications using a privacy-preserving two-phase system."
    ),
    # CLAIM 10
    (
        "A non-transitory computer-readable storage medium storing instructions that, when "
        "executed by one or more processors, cause the processors to implement:\n"
        "a. the anonymous telemedicine platform as claimed in claim 1;\n"
        "b. the automated nurse-assisted medication compliance system as claimed in claim 2;\n"
        "c. the appointment-to-consultation linking mechanism as claimed in claim 4; and\n"
        "d. the privacy-preserving notification system as claimed in claim 7;\n"
        "wherein the instructions are configured for execution on mobile computing devices "
        "communicating with a cloud-hosted backend comprising PostgreSQL with row-level security, "
        "Supabase Edge Functions on the Deno runtime, WebSocket real-time channels, Firebase "
        "Cloud Messaging, VideoSDK for voice communication, and M-Pesa mobile money integration."
    ),
]

for i, claim_text in enumerate(claims, 1):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Claim {i}.")
    run.bold = True
    doc.add_paragraph(claim_text)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("7. ABSTRACT")

doc.add_paragraph(
    "An anonymous telemedicine platform (eSIRI Plus) and method for delivering "
    "healthcare services without patient identity disclosure. The system generates "
    "unique cryptographic patient identifiers without collecting personally identifiable "
    "information. A custom HMAC-SHA256 JWT authentication system with triple-layer "
    "token storage provides session security. A multi-tier service architecture offers "
    "Economy and Royal tiers with differentiated follow-up access. An automated "
    "nurse-assisted medication compliance system enables prescribing doctors to create "
    "medication timetables; a server-side cron process detects medication times, assigns "
    "available nurses, creates voice-over-IP rooms, and dispatches nurses to call "
    "patients for medication reminders, with retry, reassignment, and text fallback "
    "mechanisms. Consultation requests use 60-second TTL with WebSocket and polling "
    "monitoring, and atomically link to pre-existing appointments via optimistic "
    "locking. Server-authoritative session timing with monotonic clock synchronization "
    "prevents client manipulation. A privacy-preserving two-phase notification system "
    "transmits only opaque identifiers in push payloads. AI-assisted medical reports "
    "are generated via large language models with QR verification codes. M-Pesa mobile "
    "money integration supports payments. The platform supports six languages and six "
    "healthcare provider categories, targeting the Tanzanian healthcare market."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. DRAWINGS AND FIGURES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("8. DRAWINGS AND FIGURES")

doc.add_paragraph(
    "The following pages contain descriptions of the drawings referenced in this "
    "specification. Formal technical drawings shall be prepared in accordance with "
    "BRELA requirements and submitted separately."
)

doc.add_paragraph()

figure_descriptions = [
    ("FIGURE 1 — System Architecture Diagram",
     "Shows the complete eSIRI Plus system architecture comprising: "
     "Patient Android Client and Doctor Android Client (Kotlin/Jetpack Compose, multi-module), "
     "connected via HTTPS REST and WebSocket to the Supabase Backend (Edge Functions on Deno, "
     "PostgreSQL with RLS, Realtime channels). External integrations shown: VideoSDK "
     "(voice/video rooms), Firebase Cloud Messaging (push notifications), M-Pesa Gateway "
     "(mobile payments), and OpenAI (report generation). Arrows indicate data flow directions: "
     "bidirectional for REST/WebSocket, unidirectional for push notifications (server→device) "
     "and payment callbacks (M-Pesa→server)."),

    ("FIGURE 2 — Patient Anonymization Flow",
     "Flowchart: START → Generate ES-XXXX-XXXX Patient ID → Create Session "
     "(UUID + 64-char token + 96-char refresh token) → Hash tokens (SHA-256 + bcrypt) → "
     "Store in patient_sessions table → Issue HMAC-SHA256 JWT (sub=session_id, role=authenticated, "
     "app_role=patient) → Prompt Security Questions (5 questions) → Hash answers (PBKDF2+salt) → "
     "Store hashed answers → Bind Device (hash device ID, store) → Session Active → END."),

    ("FIGURE 3 — Consultation Request Lifecycle",
     "Sequence diagram with actors: Patient App, Server, Doctor App. "
     "Patient App→Server: POST create request (doctor_id, service_type). "
     "Server validates doctor availability → Creates request (status=pending) → "
     "Pushes notification to Doctor App. "
     "Patient App subscribes to Realtime channel + starts 3s polling. "
     "Patient App starts 60s countdown. "
     "Doctor App→Server: POST accept/reject. "
     "If accepted: Server creates consultation, sets in_session=true, notifies Patient via Realtime. "
     "If rejected: Server notifies Patient. "
     "If 60s timeout: Patient→Server: POST expire request."),

    ("FIGURE 4 — Appointment-to-Consultation Linking",
     "Flowchart: Patient selects appointment → Validate: appointment exists? belongs to patient? "
     "consultation_id IS NULL? status in (BOOKED,CONFIRMED,MISSED)? no pending request? → "
     "Create consultation request with appointment_id → Doctor accepts → "
     "Atomic UPDATE appointments SET consultation_id=new_id, status=IN_PROGRESS "
     "WHERE consultation_id IS NULL (optimistic lock) → If update rows=0: concurrent link detected, "
     "return conflict error → If rows=1: success, navigate to consultation."),

    ("FIGURE 5 — Medication Compliance System Sequence",
     "Sequence diagram: Doctor creates timetable during report → Stored in medication_timetables. "
     "pg_cron fires every minute → medication-reminder-cron Edge Function → "
     "Convert UTC to EAT → Query matching timetables → INSERT event (ON CONFLICT DO NOTHING) → "
     "Query available nurse (specialty=nurse, verified, available, not in session) → "
     "If found: Create VideoSDK room → Push MEDICATION_REMINDER_CALL to nurse → "
     "Push MEDICATION_REMINDER_PATIENT to patient → "
     "Nurse accepts call → Voice call → Nurse ends call → "
     "POST medication-reminder-callback (completed) → Credit nurse 1000 TZS. "
     "If no nurse: Set status=no_nurse → Retry next minute (max 3). "
     "If nurse timeout (>2min): Reassign to different nurse (max 2). "
     "If all retries exhausted: Text push to patient as fallback."),

    ("FIGURE 6 — Session Management and Timer Flow",
     "State diagram: Request Accepted → ACTIVE (server sets end_time, doctor in_session=true, "
     "client syncs timer via monotonic clock) → Doctor requests extension → AWAITING_EXTENSION → "
     "Patient accepts → GRACE_PERIOD (3 min for payment) → Payment confirmed → ACTIVE (extended) → "
     "Doctor ends → COMPLETED_REPORT_PENDING → Doctor submits report → COMPLETED "
     "(trigger: in_session=false). Alternative paths: Patient declines extension → ACTIVE (original "
     "timer continues). Grace period expires without payment → COMPLETED."),

    ("FIGURE 7 — Two-Phase Notification Delivery",
     "Diagram: Phase 1: Server sends FCM with {notification_id: UUID, type: string, session_id: UUID} "
     "→ FCM delivers to device → Device wakes. "
     "Phase 2: App displays generic placeholder → App authenticates with server → "
     "App fetches full content (GET /notification/{id}) → App updates notification with actual content. "
     "For calls: Data-only FCM → IncomingCallService starts → Wake lock acquired → "
     "Full-screen call UI shown."),

    ("FIGURE 8 — AI Report Generation Flow",
     "Flowchart: Doctor submits (consultation_id, diagnosed_problem, category, severity, "
     "treatment_plan, prescriptions[]) → Server fetches service_tier → Construct LLM prompt → "
     "Call OpenAI GPT-4 → Receive 5 structured fields (symptoms, diagnosis, treatment, "
     "medications, follow_up) → Generate verification_code (UUID) → INSERT patient_reports → "
     "INSERT prescriptions → Set report_submitted=true → Trigger: in_session=false → "
     "If ROYAL + timetables: INSERT medication_timetables → Return report to doctor."),

    ("FIGURE 9 — Triple-Layer Token Storage",
     "Diagram showing three storage layers: Layer 1 (Primary): EncryptedSharedPreferences "
     "(AES-256-GCM via Android Keystore). Layer 2 (Backup): Standard SharedPreferences "
     "(fallback for Keystore failures). Layer 3 (In-Memory): TokenManager singleton cache "
     "(volatile, fastest access). "
     "Read flow: Try Layer 3 → miss → Try Layer 1 → fail → Try Layer 2 → populate Layer 3. "
     "Refresh flow: Check expiry (5-min window) → If expiring: POST refresh-patient-session → "
     "Store new tokens in all 3 layers. "
     "401 flow: TokenRefreshAuthenticator intercepts → Single refresh attempt (ReentrantLock) → "
     "Retry original request with new token."),

    ("FIGURE 10 — M-Pesa Payment Flow",
     "Sequence diagram: Patient App→Server: POST payment request (amount, phone, type) → "
     "Server→M-Pesa: STK Push initiate → M-Pesa→Patient Phone: USSD prompt → "
     "Patient enters PIN → M-Pesa processes → M-Pesa→Server: POST mpesa-callback "
     "(transaction_id, status, amount) → Server verifies callback → Updates payment record → "
     "Activates service (consultation access / call minutes / extension) → "
     "Server→Patient App: Realtime notification (payment confirmed)."),
]

for fig_title, fig_desc in figure_descriptions:
    p = doc.add_paragraph()
    run = p.add_run(fig_title)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(fig_desc)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SIGNATURE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
center_bold("DECLARATION", 14)
doc.add_paragraph()

doc.add_paragraph(
    "I, PATRICK MOKIWA KIVANDA, of Kijoleri, Ilboru Ward, Arusha, Tanzania, "
    "do hereby declare that the invention described in this specification is my own "
    "invention and that, to the best of my knowledge and belief, the facts stated "
    "in this specification are true."
)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph("Name:\tPATRICK MOKIWA KIVANDA")
p = doc.add_paragraph("National ID:\t19730427231120000127")
p = doc.add_paragraph("Date:\t" + datetime.datetime.now().strftime("%d/%m/%Y"))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph("Signature:\t____________________________________")

doc.add_paragraph()
doc.add_paragraph()

center_bold("— END OF PATENT SPECIFICATION —", 12)

# Save
doc.save("D:/esiriplus/docs/eSIRI_PLUS_PATENT_SPECIFICATION.docx")
print("Patent specification document saved to docs/eSIRI_PLUS_PATENT_SPECIFICATION.docx")
