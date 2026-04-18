from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ── Header ──
for text in [
    "United Republic of Tanzania",
    "Business Registrations and Licensing Agency",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PATENTS ACT 1987 FORM NO. 2 (Regulation 8,52)")
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Application for Registration of a Patent")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()


def add_field(label, value):
    p = doc.add_paragraph()
    run = p.add_run(label + "\t\t")
    run.bold = True
    p.add_run(str(value))


def add_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.underline = True


# ── APPLICATION ──
add_header("APPLICATION")
add_field("Tracking number", "________________________________")
add_field("Application date", datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S"))

# ── APPLICANT ──
add_header("APPLICANT")
add_field("National ID", "19730427231120000127")
add_field("Name", "PATRICK MOKIWA KIVANDA")
add_field("Gender", "Male")
add_field("Date of birth", "27/04/1973")
add_field("Nationality", "Tanzanian")
add_field("E-mail Address", "pkivanda@yahoo.com")
add_field("Mobile Phone Number", "+255784272858")

# ── TECHNICAL DESCRIPTION ──
add_header("TECHNICAL DESCRIPTION")
add_field("Patent type", "National Patent")
add_field("Title", "eSIRI PLUS: ANONYMOUS TELEMEDICINE PLATFORM WITH AUTOMATED NURSE-ASSISTED MEDICATION COMPLIANCE SYSTEM")

p = doc.add_paragraph()
run = p.add_run("Abstract")
run.bold = True

abstract = (
    "An anonymous telemedicine platform and method for providing healthcare services in "
    "resource-constrained environments without requiring patient identity disclosure. The "
    "invention comprises: a patient identity anonymization system that generates unique "
    "cryptographic patient identifiers (Patient IDs) enabling full healthcare access without "
    "collecting personally identifiable information such as names, national identification "
    "numbers, email addresses, or home addresses; a multi-tier consultation service "
    "architecture providing Economy and Royal service tiers with differentiated follow-up "
    "windows, nurse-assisted medication reminders, and tiered pricing; an automated "
    "nurse-assisted medication compliance system wherein a doctor creates a medication "
    "consumption timetable specifying medication names, dosages, administration times, and "
    "duration, and a server-side cron process automatically detects when medication times "
    "arrive, dynamically assigns an available online nurse from a pool of verified healthcare "
    "providers, creates a voice-over-IP communication channel, and dispatches the nurse to "
    "call the patient with medication context to ensure adherence to the prescribed regimen; "
    "a privacy-preserving authentication system using custom HS256 JSON Web Tokens for "
    "patients that avoids Supabase Auth dependency while maintaining session security through "
    "token rotation, proactive refresh, multi-layer storage fallback (encrypted storage, plain "
    "backup, in-memory cache), and device binding; a real-time consultation session management "
    "system with server-authoritative timers, extension workflows, grace periods for payment "
    "processing, and automatic consultation lifecycle management; a consultation request system "
    "with 60-second time-to-live requests, real-time status tracking via WebSocket subscriptions "
    "and polling fallback, and appointment-to-consultation linking enabling patients to start "
    "consultations directly from scheduled appointments with one-appointment-one-consultation "
    "enforcement; an AI-assisted consultation report generation system using large language "
    "models to produce structured medical reports with QR-code verification; and a mobile "
    "payment integration system using M-Pesa STK Push for consultation fees, video call "
    "recharges, and service access payments, all operating within a privacy-first architecture "
    "where push notification payloads contain only opaque identifiers and full content is "
    "fetched securely after device wake. The system is designed for the Tanzanian healthcare "
    "market with support for six healthcare provider categories (Nurse, Clinical Officer, "
    "Pharmacist, General Practitioner, Specialist, Psychologist) and six languages (English, "
    "Swahili, Arabic, French, Spanish, Hindi)."
)
doc.add_paragraph(abstract)

add_field("Page reference for last claims", "16")
add_field("Page reference for last description", "16")

# ── CLAIMS ──
add_header("CLAIMS")

claims = []

# CLAIM 1
claims.append((1, (
    "An anonymous telemedicine platform system, comprising:\n"
    "a patient anonymization module that generates unique cryptographic patient identifiers "
    "without collecting personally identifiable information, wherein each patient identifier "
    "follows a format of a prefix code followed by alphanumeric segments separated by hyphens;\n"
    "a multi-tier consultation service module providing at least two service tiers with "
    "differentiated follow-up access, wherein a first tier provides a single follow-up "
    "consultation and a second tier provides unlimited follow-up consultations within a "
    "configurable time window;\n"
    "a consultation request module that creates time-limited consultation requests directed "
    "to specific healthcare providers, monitors request status through real-time WebSocket "
    "subscriptions and periodic polling fallback, and links consultation requests to "
    "pre-existing appointments;\n"
    "a session management module that maintains server-authoritative consultation timers "
    "with extension workflows, grace periods, and automatic lifecycle transitions;\n"
    "a custom authentication module using cryptographic JSON Web Tokens signed with "
    "HMAC-SHA256 for patient sessions, with proactive token refresh, triple-layer token "
    "storage fallback comprising encrypted storage, plain-text backup storage, and in-memory "
    "cache, and automatic retry with token refresh on authentication failures;\n"
    "a privacy-preserving notification module wherein push notification payloads contain "
    "only opaque identifiers and notification types, and full notification content is fetched "
    "securely from a server after device activation;\n"
    "a mobile payment integration module supporting STK Push-based mobile money transactions "
    "for consultation fees and service access;\n"
    "an AI-assisted report generation module that processes consultation data through a "
    "large language model to produce structured medical reports with verification codes; and\n"
    "a multi-language localization module supporting at least six languages."
)))

# CLAIM 2
claims.append((2, (
    "The anonymous telemedicine platform system as claimed in claim 1, further comprising an "
    "automated nurse-assisted medication compliance system, wherein:\n"
    "a prescribing healthcare provider creates a medication consumption timetable for a "
    "patient through a client application, the timetable specifying at least a medication "
    "name, number of daily administrations, specific administration times in a local timezone, "
    "and a duration in days;\n"
    "the medication timetable is stored in a server database with fields including a timetable "
    "identifier, consultation reference, patient session identifier, medication name, dosage, "
    "pharmaceutical form, an array of scheduled administration times, duration, start date, "
    "computed end date, and an active status flag;\n"
    "a server-side scheduled process executes at a configurable interval, converts current "
    "server time to the local timezone, queries active timetables where the current time "
    "matches any entry in the scheduled times array and the current date falls within the "
    "start and end date range;\n"
    "for each matching timetable entry, the system inserts a reminder event record with a "
    "unique constraint on the combination of timetable identifier, scheduled date, and "
    "scheduled time to prevent duplicate processing;\n"
    "the system queries a pool of available healthcare providers filtered by specialty type, "
    "verification status, online availability status, and current session status to identify "
    "a nurse not currently engaged in another consultation;\n"
    "upon identifying an available nurse, the system creates a voice-over-IP communication "
    "room through an external video communication service API, updates the reminder event "
    "with the assigned nurse identifier, room identifier, and notification timestamp;\n"
    "the system dispatches a first push notification to the assigned nurse containing the "
    "room identifier, patient session identifier, medication name, and dosage information "
    "with a notification type indicating a medication reminder call;\n"
    "the system dispatches a second push notification to the patient indicating that a nurse "
    "will call shortly for a medication reminder;\n"
    "the nurse receives the notification on their client device, which activates an incoming "
    "call service that acquires a device wake lock and displays a full-screen call notification "
    "configured for audio-only communication;\n"
    "upon the nurse completing the call, the nurse transmits a callback to the server "
    "indicating call completion status; and\n"
    "upon successful completion, the system credits the nurse with a predetermined "
    "compensation amount of 1,000 Tanzanian Shillings with an earning type designated as "
    "medication reminder."
)))

# CLAIM 3
claims.append((3, (
    "The anonymous telemedicine platform system as claimed in claim 2, wherein the medication "
    "compliance system further comprises:\n"
    "a retry mechanism wherein reminder events with a status indicating no nurse was available "
    "are re-processed on subsequent scheduled process executions, up to a configurable "
    "maximum retry count;\n"
    "a nurse timeout mechanism wherein reminder events where the assigned nurse has been "
    "notified but has not initiated a call within a configurable timeout period are reassigned "
    "to a different nurse, excluding the previously assigned nurse, up to a configurable "
    "maximum reassignment count;\n"
    "a patient fallback mechanism wherein, after exhausting all retry and reassignment "
    "attempts, the system sends a text-only push notification directly to the patient "
    "reminding them to take the specified medication; and\n"
    "a patient-unreachable mechanism wherein the nurse can report that the patient did not "
    "answer the call, triggering a push notification to the patient indicating they missed "
    "their medication reminder."
)))

# CLAIM 4
claims.append((4, (
    "The anonymous telemedicine platform system as claimed in claim 1, wherein the "
    "consultation request module further comprises an appointment linking mechanism, wherein:\n"
    "a consultation request includes an optional appointment identifier referencing a "
    "pre-existing appointment record;\n"
    "upon receiving a consultation request with an appointment identifier, the system "
    "validates that the appointment exists, belongs to the requesting patient, has no "
    "existing consultation linked, and has a status permitting consultation initiation;\n"
    "the system further validates that no pending consultation request already exists for "
    "the same appointment to prevent duplicate requests;\n"
    "upon a healthcare provider accepting the consultation request, the system atomically "
    "updates the appointment record to link the newly created consultation and transitions "
    "the appointment status to in-progress, using an optimistic lock condition that the "
    "appointment consultation reference is still null; and\n"
    "the appointment remains visible in the patient active appointment list until a "
    "consultation is successfully linked, including appointments that have exceeded their "
    "scheduled time without a consultation."
)))

# CLAIM 5
claims.append((5, (
    "The anonymous telemedicine platform system as claimed in claim 1, wherein the patient "
    "anonymization module comprises:\n"
    "a session creation function that generates a universally unique session identifier, a "
    "cryptographically random session token of at least 64 characters, and a separate "
    "refresh token of at least 96 characters;\n"
    "a custom JSON Web Token signed using HMAC-SHA256 with a server-side secret, the token "
    "payload containing the session identifier as the subject claim, a session identifier "
    "claim, a role claim set to an authenticated role for database row-level security "
    "compatibility, an application role claim identifying the bearer as a patient, and "
    "expiration and issuance timestamps;\n"
    "the session token and refresh token are hashed using both SHA-256 for fast indexed "
    "database lookups and bcrypt for timing-attack resistant verification;\n"
    "a recovery mechanism based on five user-defined security questions whose answers are "
    "hashed using PBKDF2 before storage, with an account lockout mechanism after a "
    "configurable number of failed recovery attempts; and\n"
    "a device binding mechanism that cryptographically links a patient session to a specific "
    "device, preventing unauthorized session use across devices."
)))

# CLAIM 6
claims.append((6, (
    "The anonymous telemedicine platform system as claimed in claim 1, wherein the session "
    "management module comprises:\n"
    "a server-authoritative timer system where the consultation duration is determined by a "
    "service duration function based on healthcare provider specialty, with configurable "
    "durations per specialty type;\n"
    "a timer synchronization mechanism where client devices synchronize their countdown "
    "timers against the server time, using the difference between server-provided scheduled "
    "end time and server time to set a monotonic clock anchor, preventing manipulation "
    "through client clock changes;\n"
    "an extension workflow comprising: a healthcare provider initiating an extension request, "
    "a patient accepting or declining the extension, a grace period with configurable "
    "duration during which the patient can process payment through mobile money, and "
    "automatic consultation resumption upon payment confirmation;\n"
    "an automatic consultation lifecycle managed by database triggers that synchronize a "
    "healthcare provider in-session status based on consultation states including active, "
    "awaiting extension, grace period, and completed with report pending; and\n"
    "a follow-up system where completed consultations can be reopened within a configurable "
    "follow-up window, with the follow-up count tracked per consultation and configurable "
    "maximum follow-up limits based on service tier."
)))

# CLAIM 7
claims.append((7, (
    "The anonymous telemedicine platform system as claimed in claim 1, wherein the "
    "privacy-preserving notification module comprises:\n"
    "a two-phase notification delivery system where the first phase transmits a push "
    "notification containing only an opaque notification identifier and a notification type "
    "string, with no patient health information, provider names, or consultation details "
    "included in the push payload;\n"
    "upon device activation from the push notification, a second phase wherein the client "
    "application authenticates with the server using stored session credentials and fetches "
    "the full notification content including title, body text, and associated metadata;\n"
    "a generic notification display that immediately shows a placeholder notification with "
    "a type-appropriate generic title while the secure content fetch is in progress, "
    "subsequently updating the notification with the actual content upon successful retrieval;\n"
    "specialized notification handling for incoming call notifications where the push payload "
    "is configured as a data-only message without a display notification block, enabling a "
    "foreground service to acquire a wake lock and display a full-screen call notification "
    "on devices with aggressive background process management; and\n"
    "FCM token synchronization where new Firebase Cloud Messaging tokens are automatically "
    "synchronized with the server, with separate token storage for healthcare provider "
    "accounts and patient session accounts."
)))

# CLAIM 8
claims.append((8, (
    "The anonymous telemedicine platform system as claimed in claim 1, wherein the "
    "AI-assisted report generation module comprises:\n"
    "a report request interface accepting at minimum a consultation identifier, diagnosed "
    "problem, medical category, severity assessment, treatment plan, optional prescriptions "
    "with medication name, pharmaceutical form, dosage instructions, and route of "
    "administration;\n"
    "a large language model integration that generates five structured text fields from the "
    "input: presenting symptoms, diagnosis assessment, treatment plan prose incorporating "
    "prescribed medications, prescribed medications prose with formal dosage instructions, "
    "and follow-up instructions;\n"
    "a verification code generated as a universally unique identifier and embedded in the "
    "report for authenticity verification;\n"
    "a report submission trigger that sets a report-submitted flag on the consultation "
    "record, which activates a database trigger function to clear the healthcare provider "
    "in-session status, making them available for new consultation requests;\n"
    "for consultations of the premium service tier, the report module accepts optional "
    "medication timetable specifications and creates corresponding database records for the "
    "automated nurse-assisted medication compliance system; and\n"
    "structured prescription storage in both a JSON column on the report record for complete "
    "preservation and individual prescription table rows for relational queries."
)))

# CLAIM 9
claims.append((9, (
    "A method for providing anonymous telemedicine services with automated medication "
    "compliance, comprising the steps of:\n"
    "generating a unique cryptographic patient identifier without collecting personally "
    "identifiable information;\n"
    "establishing a patient session using a custom HMAC-SHA256 signed JSON Web Token with "
    "triple-layer token storage;\n"
    "receiving a consultation request from the patient directed to a specific healthcare "
    "provider, the request optionally linked to a pre-existing appointment;\n"
    "monitoring the consultation request status through concurrent real-time WebSocket "
    "subscription and periodic server polling;\n"
    "upon provider acceptance, creating a consultation session with a server-authoritative "
    "timer synchronized to the client through monotonic clock anchoring;\n"
    "managing the consultation lifecycle through extension requests, grace periods, and "
    "payment integration;\n"
    "upon consultation completion, generating an AI-assisted medical report through a large "
    "language model;\n"
    "for premium-tier consultations, receiving medication timetable specifications from the "
    "healthcare provider and storing them in a database;\n"
    "executing a scheduled server process at configurable intervals to detect medication "
    "administration times;\n"
    "for each detected medication time, querying available nurses, creating a voice "
    "communication channel, and dispatching call notifications;\n"
    "upon nurse call completion, crediting the nurse earnings account; and\n"
    "transmitting all push notifications using a privacy-preserving two-phase delivery system."
)))

# CLAIM 10
claims.append((10, (
    "A non-transitory computer-readable storage medium storing instructions that, when "
    "executed by one or more processors, cause the one or more processors to implement:\n"
    "an anonymous telemedicine platform as claimed in claim 1;\n"
    "an automated nurse-assisted medication compliance system as claimed in claim 2;\n"
    "an appointment-to-consultation linking mechanism as claimed in claim 4; and\n"
    "a privacy-preserving notification system as claimed in claim 7;\n"
    "wherein the instructions are configured for execution on mobile computing devices "
    "running the Android operating system with a Kotlin-based client application "
    "communicating with a cloud-hosted backend comprising PostgreSQL database with row-level "
    "security policies, Supabase Edge Functions implemented in TypeScript running on the "
    "Deno runtime, real-time event delivery through WebSocket channels, Firebase Cloud "
    "Messaging for push notification delivery, VideoSDK for voice and video communication, "
    "and M-Pesa mobile money integration for payment processing."
)))

for num, desc in claims:
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"CLAIM {num}")
    run.bold = True
    run.font.size = Pt(12)
    add_field("Claim number", str(num))
    p = doc.add_paragraph()
    run = p.add_run("Claim description")
    run.bold = True
    doc.add_paragraph(desc)

# ── OWNERS ──
doc.add_page_break()
add_header("OWNERS")
p = doc.add_paragraph()
run = p.add_run("OWNER 1")
run.bold = True
add_field("National ID", "19730427231120000127")
add_field("Name", "PATRICK MOKIWA KIVANDA")
add_field("Gender", "Male")
add_field("Date of birth", "27/04/1973")
add_field("Nationality", "Tanzanian")
add_field("E-mail Address", "pkivanda@yahoo.com")
add_field("Mobile Phone Number", "0784272858")
add_field("Residential address",
    "Tanzania, Region Arusha, District Arusha, Ward Ilboru, "
    "Postal code 23227, Street Kijoleri, Road TAKUKURLI RD, "
    "Plot number 24, Block number A 202, House number 3.")

# ── REPRESENTATIVES ──
add_header("REPRESENTATIVES")
p = doc.add_paragraph()
run = p.add_run("REPRESENTATIVE 1")
run.bold = True
doc.add_paragraph()

# ── PRIORITIES ──
add_header("PRIORITIES")
p = doc.add_paragraph()
run = p.add_run("PRIORITY 1")
run.bold = True
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph("PATRICK MOKIWA KIVANDA")
p = doc.add_paragraph()
p.add_run("\t\t\t\t\t____________________")
p = doc.add_paragraph()
p.add_run("\t\t\t\t\tSignature and date")

doc.save("D:/esiriplus/docs/eSIRI_PLUS_PATENT_APPLICATION.docx")
print("Patent document saved successfully!")
